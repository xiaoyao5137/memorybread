#!/usr/bin/env python3
"""Generate the polished DOCX delivery guide from its governed Markdown source."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PRESET = "compact_reference_guide"
HEADER_PATTERN = "memo_masthead"
PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM_DXA = 80
CELL_MARGIN_START_END_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CODE_FILL = "F7F4EF"
CODE_BORDER = "D8CFC3"
WHITE = "FFFFFF"
BLACK = "111111"
BODY_FONT = "Calibri"
# Named macOS override for Chinese glyph coverage while preserving the preset's
# Calibri Latin typography.
CJK_FONT = "Hiragino Sans GB"
MONO_FONT = "Menlo"


def set_run_font(run, name=BODY_FONT, east_asia=CJK_FONT, size=None, color=None,
                 bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_MARGIN_TOP_BOTTOM_DXA,
                     bottom=CELL_MARGIN_TOP_BOTTOM_DXA,
                     start=CELL_MARGIN_START_END_DXA,
                     end=CELL_MARGIN_START_END_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C2CC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440.0)
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_paragraph_bottom_border(paragraph, color=BLUE, size="14", space="7"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    set_run_font(run, size=9, color=MUTED)


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(3)
    left = paragraph.add_run("记忆面包  |  应用内更新与版本管理")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    add_paragraph_bottom_border(paragraph, color="D5DCE4", size="6", space="4")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(3)
    label = p.add_run("实施交付  ·  2026-08-05  ·  第 ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(p)
    suffix = p.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for name in ("MB Bullet", "MB Number"):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.base_style = normal
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "MB Code" not in styles:
        code = styles.add_style("MB Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["MB Code"]
    code.font.name = MONO_FONT
    code.font.size = Pt(8.2)
    code.font.color.rgb = RGBColor.from_string(INK)
    code._element.rPr.rFonts.set(qn("w:ascii"), MONO_FONT)
    code._element.rPr.rFonts.set(qn("w:hAnsi"), MONO_FONT)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.08

    if "Table Citation" not in styles:
        citation = styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        citation = styles["Table Citation"]
    citation.base_style = normal
    citation.font.size = Pt(9)
    citation.font.color.rgb = RGBColor.from_string(MUTED)
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element

    def next_id(tag, attr):
        values = []
        for node in numbering.findall(qn(tag)):
            raw = node.get(qn(attr))
            if raw is not None and raw.isdigit():
                values.append(int(raw))
        return max(values or [0]) + 1

    def add_definition(marker, number_format):
        abstract_id = next_id("w:abstractNum", "w:abstractNumId")
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), number_format)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        level.append(lvl_text)
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        level.append(justification)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "540")
        indent.set(qn("w:hanging"), "271")
        p_pr.append(indent)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        abstract.append(level)
        numbering.append(abstract)

        num_id = next_id("w:num", "w:numId")
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
        return num_id

    return add_definition("•", "bullet"), add_definition("%1.", "decimal")


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))


def restart_numbering(doc, source_num_id, start_value):
    numbering = doc.part.numbering_part.element
    source = None
    ids = []
    for node in numbering.findall(qn("w:num")):
        raw = node.get(qn("w:numId"))
        if raw is not None and raw.isdigit():
            ids.append(int(raw))
        if raw == str(source_num_id):
            source = node
    assert source is not None
    source_abstract = source.find(qn("w:abstractNumId"))
    assert source_abstract is not None

    num_id = max(ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), source_abstract.get(qn("w:val")))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), str(start_value))
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), BODY_FONT)
    r_fonts.set(qn("w:hAnsi"), BODY_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.extend((r_fonts, color, underline))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text, size=None, color=None):
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=MONO_FONT, east_asia=CJK_FONT,
                         size=(size or 10.2), color=INK)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), CODE_FILL)
            run._element.get_or_add_rPr().append(shading)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def add_title_page(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("MEMORYBREAD  ·  TECHNICAL DELIVERY")
    set_run_font(run, size=9.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("应用内更新与版本管理")
    set_run_font(run, size=25, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("全应用签名更新、双发行轨道、灰度发布与运营管理完整方案")
    set_run_font(run, size=13.5, color=MUTED)

    metadata = (
        ("状态", "已实现；生产发布凭证与线上对象地址待配置"),
        ("适用范围", "MemoryBread · mb-admin · mb-ops"),
        ("发布边界", "官网直装版应用内安装 / Mac App Store 版商店更新"),
        ("最后更新", "2026-08-05"),
    )
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.1
        label_run = paragraph.add_run(label + "：")
        set_run_font(label_run, size=10.5, color=BLACK, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=BLACK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(7)
    rule.paragraph_format.space_after = Pt(10)
    add_paragraph_bottom_border(rule, color=BLUE, size="18", space="7")

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [CONTENT_WIDTH_DXA])
    set_table_borders(callout, color="D7E1EB", size="5")
    set_repeat_table_header(callout.rows[0])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_inline(
        paragraph,
        "交付结论：直装版已具备应用内下载、签名与 SHA-256 双重校验、完整应用安装和重启；App Store 版保留合规商店更新路径。",
        size=10.5,
        color=INK,
    )


def split_table_row(line):
    value = line.strip().strip("|")
    return [cell.strip() for cell in re.split(r"(?<!\\)\|", value)]


def is_separator_row(cells):
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def text_weight(value):
    score = 0
    for char in re.sub(r"[*`]", "", value):
        score += 2 if ord(char) > 127 else 1
    return max(5, min(score, 38))


def content_fit_widths(rows):
    column_count = len(rows[0])
    if column_count == 1:
        return [CONTENT_WIDTH_DXA]
    if column_count == 2:
        first = max(text_weight(row[0]) for row in rows)
        second = max(text_weight(row[1]) for row in rows)
        if first <= 18 and second >= first * 2:
            return [2700, 6660]
    minimum = 1050 if column_count <= 4 else 850
    weights = [max(text_weight(row[index]) for row in rows) for index in range(column_count)]
    available = CONTENT_WIDTH_DXA - minimum * column_count
    total = sum(weights)
    widths = [minimum + int(available * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    widths = content_fit_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
                add_inline(paragraph, value, size=9.1, color=INK)
                for run in paragraph.runs:
                    run.bold = True
            else:
                add_inline(paragraph, value, size=8.9, color=BLACK)
    after = doc.add_paragraph(style="Table Citation")
    after.add_run("")


def add_code_block(doc, language, lines):
    label = doc.add_paragraph(style="Table Citation")
    label_run = label.add_run((language or "text").upper() + "  ·  配置/流程示例")
    set_run_font(label_run, size=8.5, color=MUTED, bold=True)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=CODE_BORDER, size="5")
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_FILL)
    set_cell_margins(cell, top=100, bottom=100)
    first = cell.paragraphs[0]
    first.style = "MB Code"
    for index, line in enumerate(lines or [""]):
        paragraph = first if index == 0 else cell.add_paragraph(style="MB Code")
        run = paragraph.add_run(line.replace("\t", "    "))
        set_run_font(run, name=MONO_FONT, east_asia=CJK_FONT, size=8.2, color=INK)
    after = doc.add_paragraph(style="Table Citation")
    after.add_run("")


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D7E1EB", size="5")
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_inline(paragraph, text, size=10.2, color=INK)


def add_markdown_body(doc, markdown, bullet_num_id, decimal_num_id):
    lines = markdown.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 摘要"))
    index = start
    active_decimal_num_id = decimal_num_id
    last_source_number = None
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, language, code_lines)
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines):
            header = split_table_row(stripped)
            separator = split_table_row(lines[index + 1])
            if len(header) == len(separator) and is_separator_row(separator):
                rows = [header]
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    row = split_table_row(lines[index])
                    if len(row) == len(header):
                        rows.append(row)
                    index += 1
                add_markdown_table(doc, rows)
                continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            last_source_number = None
            level = len(heading.group(1)) - 1
            paragraph = doc.add_paragraph(style="Heading " + str(level))
            add_inline(paragraph, heading.group(2))
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            paragraph = doc.add_paragraph(style="MB Bullet")
            apply_numbering(paragraph, bullet_num_id)
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            source_number = int(numbered.group(1))
            if last_source_number is None or source_number != last_source_number + 1:
                active_decimal_num_id = restart_numbering(
                    doc, decimal_num_id, source_number
                )
            paragraph = doc.add_paragraph(style="MB Number")
            apply_numbering(paragraph, active_decimal_num_id)
            add_inline(paragraph, numbered.group(2))
            last_source_number = source_number
            index += 1
            continue
        if stripped.startswith(">"):
            add_callout(doc, stripped.lstrip("> "))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (candidate.startswith(("#", "- ", ">", "```", "|"))
                    or re.match(r"^\d+\.\s+", candidate)):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))


def audit_document(doc):
    section = doc.sections[0]
    assert round(section.page_width.inches, 3) == 8.5
    assert round(section.page_height.inches, 3) == 11.0
    for margin in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin):
        assert round(margin.inches, 3) == 1.0
    assert round(section.header_distance.inches, 3) == 0.492
    assert round(section.footer_distance.inches, 3) == 0.492
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == str(CONTENT_WIDTH_DXA)
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == str(TABLE_INDENT_DXA)


def main():
    root = Path(__file__).resolve().parents[1]
    source = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else root / "doc" / "应用内更新与版本管理完整方案.md"
    output = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else source.with_suffix(".docx")
    markdown = source.read_text(encoding="utf-8")

    doc = Document()
    configure_page(doc.sections[0])
    configure_header_footer(doc.sections[0])
    configure_styles(doc)
    bullet_num_id, decimal_num_id = create_numbering(doc)
    doc.core_properties.title = "记忆面包应用内更新与版本管理完整方案"
    doc.core_properties.subject = "全应用签名更新、版本管理、灰度发布与运营流程"
    doc.core_properties.author = "MemoryBread"
    doc.core_properties.keywords = "MemoryBread, Tauri, updater, version management"

    add_title_page(doc)
    add_markdown_body(doc, markdown, bullet_num_id, decimal_num_id)
    audit_document(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print("Generated {} using preset={} header={}".format(output, PRESET, HEADER_PATTERN))


if __name__ == "__main__":
    main()
